"""
Document Parser
Supports multiple document formats: TXT, PDF, DOCX, MD
"""

import os
from typing import List, Optional
from pathlib import Path
import logging

# Initialize logger
logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents from various formats"""
    
    SUPPORTED_FORMATS = ['txt', 'pdf', 'docx', 'md']
    
    def __init__(self):
        """Initialize document parser"""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are available"""
        self.has_pypdf = False
        self.has_docx = False
        
        try:
            import pypdf
            self.has_pypdf = True
        except ImportError:
            logger.warning("pypdf not available. PDF parsing will be disabled.")
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not available. DOCX parsing will be disabled.")
    
    def parse_file(self, filepath: str) -> str:
        """
        Parse a document file and extract text
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Extracted text content
        """
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")
        
        extension = filepath.suffix.lower().lstrip('.')
        
        if extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {extension}. Supported: {self.SUPPORTED_FORMATS}")
        
        # Route to appropriate parser
        if extension == 'txt' or extension == 'md':
            return self._parse_text(filepath)
        elif extension == 'pdf':
            return self._parse_pdf(filepath)
        elif extension == 'docx':
            return self._parse_docx(filepath)
        else:
            raise ValueError(f"No parser available for {extension}")
    
    def _parse_text(self, filepath: Path) -> str:
        """Parse plain text or markdown file"""
        # Try common encodings in order
        encodings = ['utf-8', 'utf-16', 'gbk', 'gb2312', 'latin-1', 'ascii']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        
        # If all encodings fail, read as binary and decode with errors='replace'
        with open(filepath, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')
    
    def _parse_pdf(self, filepath: Path) -> str:
        """Parse PDF file"""
        if not self.has_pypdf:
            raise RuntimeError("pypdf not installed. Install with: pip install pypdf")
        
        import pypdf
        
        text_content = []
        
        try:
            with open(filepath, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
            
            return '\n\n'.join(text_content)
        
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {e}")
    
    def _parse_docx(self, filepath: Path) -> str:
        """Parse DOCX file"""
        if not self.has_docx:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        
        import docx
        
        try:
            doc = docx.Document(filepath)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n\n'.join(text_content)
        
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {e}")
    
    def parse_directory(self, directory: str, recursive: bool = True) -> List[dict]:
        """
        Parse all supported documents in a directory
        
        Args:
            directory: Directory path
            recursive: Whether to search recursively
            
        Returns:
            List of dicts with 'filepath' and 'content' keys
        """
        directory = Path(directory)
        
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Invalid directory: {directory}")
        
        documents = []
        
        # Find all supported files
        if recursive:
            files = []
            for ext in self.SUPPORTED_FORMATS:
                files.extend(directory.rglob(f"*.{ext}"))
        else:
            files = []
            for ext in self.SUPPORTED_FORMATS:
                files.extend(directory.glob(f"*.{ext}"))
        
        # Parse each file
        for filepath in files:
            try:
                content = self.parse_file(str(filepath))
                documents.append({
                    'filepath': str(filepath),
                    'filename': filepath.name,
                    'content': content
                })
                logger.info(f"Successfully parsed: {filepath.name}")
            except Exception as e:
                logger.error(f"Failed to parse {filepath.name}: {e}")
        
        return documents
